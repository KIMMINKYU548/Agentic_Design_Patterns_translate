from pathlib import Path
from docxcompose.composer import Composer
from docx import Document
import mammoth
from bs4 import BeautifulSoup

from cfg import MAMMOTH_STYLE_MAP

def merge_docx_in_order(file_list: list[Path], master_docx: Path,
                        cover_docx: Path | None = None,
                        toc_docx: Path | None = None) -> tuple[Path, bool]:
    base = Document()
    comp = Composer(base)

    # 1) 표지
    if cover_docx and cover_docx.exists():
        comp.append(Document(str(cover_docx)))
    else:
        # 심플 표지
        c = Document()
        c.add_heading("Agentic Design Patterns", 0)
        c.add_paragraph("Korean Edition (Auto-compiled)")
        tmp = master_docx.parent / "_auto_cover.docx"
        c.save(tmp); comp.append(Document(str(tmp)))

    # 2) 목차(원문을 docx로 제공한 경우)
    if toc_docx and toc_docx.exists():
        comp.append(Document(str(toc_docx)))
        insert_auto_toc = False
    else:
        insert_auto_toc = True  # HTML 단계에서 자동 TOC 생성

    # 3) 본문
    for f in file_list:
        comp.append(Document(str(f)))

    comp.save(str(master_docx))
    return master_docx, insert_auto_toc

def master_docx_to_html(master_docx: Path, master_html: Path, insert_auto_toc: bool) -> Path:
    with open(master_docx, "rb") as f:
        r = mammoth.convert_to_html(f, style_map=MAMMOTH_STYLE_MAP)
    soup = BeautifulSoup(r.value, "lxml")

    if insert_auto_toc:
        heads = soup.find_all(["h1","h2"])
        if heads:
            toc = soup.new_tag("div", **{"class":"auto-toc"})
            toc_h = soup.new_tag("h1"); toc_h.string = "Contents"
            toc.append(toc_h)
            ol = soup.new_tag("ol")
            for i,h in enumerate(heads, start=1):
                hid = h.get("id") or f"h{i}"
                h["id"] = hid
                li = soup.new_tag("li")
                a = soup.new_tag("a", href=f"#{hid}")
                a.string = h.get_text(strip=True)[:200]
                li.append(a); ol.append(li)
            toc.append(ol)
            body = soup.body or soup
            body.insert(0, toc)

    master_html.write_text(str(soup), encoding="utf-8")
    return master_html

